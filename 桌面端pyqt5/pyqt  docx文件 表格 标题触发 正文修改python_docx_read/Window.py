from PyQt5.QtWidgets import QWidget,QFileDialog,QTableWidgetItem, QTableWidget, QApplication, QPushButton, QMenu ,QLabel ,QLineEdit,QMessageBox, QTextEdit, QVBoxLayout
import sys
import os
import docx
from docx import Document


class Example(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
        self.file_name = ""

    def initUI(self):
        self.resize(700,700)
        self.setWindowTitle('DOCX助手')
       # self.count = 10 #默认10s

        self.count_num = 0

        self.text1 = QLabel("未指定", self)
        self.text1.resize(500, 50)
        self.text1.move(10, 10)
        self.text2 = QLabel("Table:"+"1", self)
        self.text2.move(0, 600)

        self.bt1 = QPushButton("打开文件",self)
        self.bt1.resize(100, 50)
        self.bt1.move(50, 50)
        self.bt2 = QPushButton("上一个图", self)
        self.bt2.resize(100, 100)
        self.bt2.move(150, 500)
        self.bt3 = QPushButton("下一个图", self)
        self.bt3.resize(100, 100)
        self.bt3.move(150,600)
        self.bt4 = QPushButton("保存图表", self)
        self.bt4.resize(100, 100)
        self.bt4.move(400, 500)
        self.bt5 = QPushButton("保存各级标题", self)
        self.bt5.resize(100, 100)
        self.bt5.move(400, 600)

        self.textEdit = QTextEdit(self)
        self.textEdit.resize(500, 400)
        self.textEdit.move(100, 100)

        # self.textEdit.show()
        #点击文件
        self.click1()
        ############################################################################################
        self.bt1.clicked.connect(self.click1)
        self.bt2.clicked.connect(self.click2)
        self.bt3.clicked.connect(self.click3)
        self.bt4.clicked.connect(self.click4)
        self.bt5.clicked.connect(self.click5)

        # self.show()




###############################################################
    #打开
    def click1(self):
        self.file_name = QFileDialog.getOpenFileName(self)[0]
        self.count_num = 0
        self.textEdit.setText("")
        self.out_file_name = self.file_name
        self.text1.setText(self.file_name)
        self.obj = Document(self.file_name)
        self.tables = self.obj.tables  # 获取文件中的表格集
        self.paragraphs = []
        for p in self.obj.paragraphs:
            # p = obj.paragraphs[ii]
            style_name = p.style.name
            if style_name.startswith('Heading'):
                #add into
                self.paragraphs.append(p)
                # print(style_name, p.text, sep=':')
                # #删除
                # p.clear()
                # #添加
                # p.add_run("added")
                self.textEdit.append(p.text)
        ####
        self.show_temp_table()

    #显示table
    def show_temp_table(self):
        self.row_count = 0
        self.col_count = self.tables[self.count_num]._column_count
        self.data = []
        for i, row in enumerate(self.tables[self.count_num].rows[:]):  # 读每行
            self.row_count += 1
            row_content = []
            for cell in row.cells[:]:  # 读一行中的所有单元格
                c = cell.text
                # 修改数据
                # cell.text = "231"
                row_content.append(c)
            # print(row_content)  # 以列表形式导出每一行数据
            self.data.append(row_content)
        self.tableWidget = QTableWidget(self.row_count,self.col_count)
        self.tableWidget.resize(500,500)
        self.tableWidget.move(0,0)
        for ii in range(self.row_count):
            for jj in range(self.col_count):
                newItem = QTableWidgetItem(self.data[ii][jj])
                self.tableWidget.setItem(ii,jj,newItem)
        self.tableWidget.show()

    #last
    def click2(self):
        self.count_num-=1
        if self.count_num<0:
            self.count_num = 0
        self.text2.setText("Table:"+str(self.count_num+1))
        self.show_temp_table()
    #next
    def click3(self):
        self.count_num+=1
        if self.count_num>=len(self.tables):
            self.count_num = len(self.tables) - 1
        self.text2.setText("Table:"+str(self.count_num+1))
        self.show_temp_table()

    #保存图表
    def click4(self):
        # print(self.tableWidget.item(1,1).text())
        for ii in range(self.row_count):
            for jj in range(self.col_count):
                if self.tableWidget.item(ii,jj).text()!= self.data[ii][jj]:
                    print("changed")
                    self.tables[self.count_num].cell(ii,jj).text = self.tableWidget.item(ii,jj).text()
        self.obj.save(self.out_file_name)
        QMessageBox.warning(self, "提示", "保存完成")

    #保存head
    def click5(self):
        #setPlainText\append
        changed_txt = self.textEdit.toPlainText()
        head_list = changed_txt.split("\n")
        for ii in range(len(head_list)):
            if self.paragraphs[ii].text.strip() != head_list[ii].strip():
                self.paragraphs[ii].clear()
                self.paragraphs[ii].add_run(head_list[ii])

        self.obj.save(self.out_file_name)
        QMessageBox.warning(self, "提示", "保存完成")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Example()
    ex.show()
    sys.exit(app.exec_())